import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn #, nsdecls
# from docx.oxml import parse_xml # Not used directly, can be removed if not needed elsewhere
from docx.oxml.shared import OxmlElement
import os
import re
import json
from PyQt5.QtCore import QUrl # For converting local paths to file URLs

class MarkdownToWordConverter:
    DEFAULT_STYLES_CONFIG = {
        'H1': {
            'font_name': '黑体', 'font_size': 24, 'bold': True, 'italic': False, 'color_rgb': (0, 0, 0),
            'space_before_pt': 12, 'space_after_pt': 6, 'alignment': 'LEFT'
        },
        'H2': {
            'font_name': '黑体', 'font_size': 20, 'bold': True, 'italic': False, 'color_rgb': (0, 0, 0),
            'space_before_pt': 10, 'space_after_pt': 5, 'alignment': 'LEFT'
        },
        'H3': {
            'font_name': '黑体', 'font_size': 18, 'bold': True, 'italic': False, 'color_rgb': (0, 0, 0),
            'space_before_pt': 8, 'space_after_pt': 4, 'alignment': 'LEFT'
        },
        'H4': {
            'font_name': '宋体', 'font_size': 16, 'bold': True, 'italic': False, 'color_rgb': (50, 50, 50),
            'space_before_pt': 6, 'space_after_pt': 3, 'alignment': 'LEFT'
        },
        'H5': {
            'font_name': '宋体', 'font_size': 14, 'bold': True, 'italic': False, 'color_rgb': (70, 70, 70),
            'space_before_pt': 5, 'space_after_pt': 2, 'alignment': 'LEFT'
        },
        'H6': {
            'font_name': '宋体', 'font_size': 12, 'bold': True, 'italic': False, 'color_rgb': (90, 90, 90),
            'space_before_pt': 4, 'space_after_pt': 2, 'alignment': 'LEFT'
        },
        'paragraph': {
            'font_name': '宋体', 'font_size': 12, 'bold': False, 'italic': False, 'color_rgb': (0, 0, 0),
            'line_spacing': 1.5, 'first_line_indent_cm': 0.0, # Default to no indent, can be set e.g. 0.74
            'space_before_pt': 0, 'space_after_pt': 6, 'alignment': 'LEFT' # Default 6pt after paragraph
        },
        'code_block': {
            'font_name': 'Consolas', 'font_size': 10, 'bold': False, 'italic': False, 'color_rgb': (0, 0, 0),
            'background_color': "F0F0F0", 'line_spacing': 1.0, 'space_before_pt': 6, 'space_after_pt': 6
        },
        'inline_code': { # Style for inline <code> elements
            'font_name': 'Consolas', # Will use this font
            'font_size_ratio': 0.9,  # Relative to surrounding paragraph font size
            'color_rgb': (50, 50, 50),
            'background_color': "EEEEEE" # Optional background for inline code in Word (via run shading)
        }
    }

    def __init__(self, styles_config=None):
        # Make a deep copy to prevent modification of class default or shared instances
        if styles_config is not None:
            self.styles_config = {k: v.copy() for k, v in styles_config.items()}
        else:
            self.styles_config = {k: v.copy() for k, v in self.DEFAULT_STYLES_CONFIG.items()}
        self._ensure_rgb_tuples()


    def _ensure_rgb_tuples(self):
        """Ensure all color_rgb values are tuples, as JSON loads them as lists."""
        for style_name, config in self.styles_config.items():
            if 'color_rgb' in config and isinstance(config['color_rgb'], list):
                config['color_rgb'] = tuple(config['color_rgb'])

    def _apply_font_style(self, run, style_config, base_font_size_pt=None):
        font = run.font
        font.name = style_config.get('font_name', 'Calibri') # Default to Calibri if not specified
        # For East Asian fonts, it's crucial to set w:eastAsia
        if font.name: # Check if font_name is not empty
             font.element.rPr.rFonts.set(qn('w:eastAsia'), font.name)

        font_size = style_config.get('font_size')
        font_size_ratio = style_config.get('font_size_ratio')

        if font_size is not None:
            font.size = Pt(font_size)
        elif font_size_ratio is not None and base_font_size_pt is not None:
            font.size = Pt(base_font_size_pt * font_size_ratio)
        else:
            font.size = Pt(11) # Default font size if none specified

        font.bold = style_config.get('bold', False)
        font.italic = style_config.get('italic', False)
        
        color_rgb = style_config.get('color_rgb')
        if color_rgb and isinstance(color_rgb, (tuple, list)) and len(color_rgb) == 3:
            font.color.rgb = RGBColor(*color_rgb)
        
        # Background color for inline code (applied as run shading)
        if 'background_color' in style_config and style_config.get('background_color'):
            # This is for inline code, code blocks use cell background
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), style_config['background_color'])
            run.element.rPr.append(shd)


    def _apply_paragraph_format(self, paragraph, style_config):
        p_format = paragraph.paragraph_format
        if 'space_before_pt' in style_config:
            p_format.space_before = Pt(style_config['space_before_pt'])
        if 'space_after_pt' in style_config:
            p_format.space_after = Pt(style_config['space_after_pt'])
        if 'line_spacing' in style_config:
            p_format.line_spacing = style_config['line_spacing']
        if 'first_line_indent_cm' in style_config and style_config['first_line_indent_cm'] > 0:
            p_format.first_line_indent = Cm(style_config['first_line_indent_cm'])
        else:
            p_format.first_line_indent = None # Remove indent if 0 or not specified

        alignment_map = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        p_format.alignment = alignment_map.get(str(style_config.get('alignment', 'LEFT')).upper(), WD_ALIGN_PARAGRAPH.LEFT)


    def _set_cell_background(self, cell, hex_color_string):
        if hex_color_string:
            if hex_color_string.startswith('#'):
                hex_color_string = hex_color_string[1:]
            if len(hex_color_string) == 6: # Ensure it's a valid hex
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), hex_color_string.upper())
                # Ensure tcPr exists
                if cell._tc.tcPr is None:
                    cell._tc.get_or_add_tcPr()
                cell._tc.tcPr.append(shd)
            
    def _add_image_to_paragraph(self, paragraph, img_element, md_dir):
        img_src = img_element.get('src')
        alt_text = img_element.get('alt', '')
        
        # Resolve image path (can be web URL or local path)
        if img_src.startswith(('http://', 'https://')):
            # For web URLs, python-docx cannot directly embed. User needs to download and insert.
            # For now, we'll just add a placeholder text.
            run = paragraph.add_run(f"[Web Image: {alt_text or img_src}]")
            run.italic = True
            print(f"Info: Web image '{img_src}' cannot be embedded directly. Added placeholder.")
            return

        # Handle local file paths, including file:/// URLs
        local_img_path = ""
        if img_src.startswith('file:///'):
            local_img_path = QUrl(img_src).toLocalFile()
        elif not os.path.isabs(img_src) and md_dir: # Relative path
            local_img_path = os.path.join(md_dir, img_src)
        elif os.path.isabs(img_src): # Absolute path
            local_img_path = img_src
        
        if local_img_path and os.path.exists(local_img_path):
            try:
                # Attempt to add picture. Consider desired width.
                # For now, default width. Can be enhanced to parse width from style or attributes.
                paragraph.add_run().add_picture(local_img_path) # Default width
            except Exception as e:
                print(f"Warning: Could not insert image '{local_img_path}': {e}")
                paragraph.add_run(f"[Image load error: {alt_text or os.path.basename(local_img_path)}]")
        else:
            print(f"Warning: Image file not found or path invalid: '{img_src}' (resolved to '{local_img_path}')")
            paragraph.add_run(f"[Image not found: {alt_text or img_src}]")

    def _add_inline_content_to_paragraph(self, paragraph, content_element, parent_style_config, md_dir):
        """
        Adds inline content (text, strong, em, code, img, a) to a paragraph.
        Handles recursive processing of inline children.
        """
        base_font_size = parent_style_config.get('font_size', 12) # Get base font size from parent (paragraph)

        if isinstance(content_element, NavigableString):
            text = str(content_element)
            # Markdown often adds newlines that result in empty strings after strip, ignore them unless they are actual spaces.
            if text.strip() or text == " ": # Keep single spaces
                run = paragraph.add_run(text)
                self._apply_font_style(run, parent_style_config, base_font_size)
        elif isinstance(content_element, Tag):
            tag_name = content_element.name
            # print(f"Processing inline tag: {tag_name}")

            if tag_name in ['strong', 'b']:
                # Create a new style config for bold, inheriting from parent
                current_run_style = parent_style_config.copy()
                current_run_style['bold'] = True
                for child in content_element.contents:
                    self._add_inline_content_to_paragraph(paragraph, child, current_run_style, md_dir)
            elif tag_name in ['em', 'i']:
                current_run_style = parent_style_config.copy()
                current_run_style['italic'] = True
                for child in content_element.contents:
                    self._add_inline_content_to_paragraph(paragraph, child, current_run_style, md_dir)
            elif tag_name == 'code':
                # Use 'inline_code' style, but pass paragraph's font size for ratio calculation
                inline_code_style_config = self.styles_config.get('inline_code', {}).copy()
                # Text for inline code
                code_text = content_element.get_text()
                if code_text:
                    run = paragraph.add_run(code_text)
                    self._apply_font_style(run, inline_code_style_config, base_font_size_pt=base_font_size)
            elif tag_name == 'img':
                self._add_image_to_paragraph(paragraph, content_element, md_dir)
            elif tag_name == 'a':
                # For links, for now, just add text with parent style. Hyperlinking is more complex.
                link_text = content_element.get_text()
                # href = content_element.get('href', '')
                run = paragraph.add_run(link_text)
                self._apply_font_style(run, parent_style_config, base_font_size)
                # Could add hyperlink styling (blue, underline) here if desired, even without full hyperlink functionality
                # run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                # run.font.underline = True
            elif tag_name == 'br':
                paragraph.add_run().add_break()
            else: # Other inline tags (e.g., span, sub, sup - though not standard MD)
                for child in content_element.children:
                    self._add_inline_content_to_paragraph(paragraph, child, parent_style_config, md_dir)


    def _process_html_element(self, doc, element, md_dir, current_list_level=0, in_list_item=False):
        tag_name = element.name
        # print(f"Processing block tag: {tag_name}, list_level: {current_list_level}, in_list_item: {in_list_item}")

        para_style_config = self.styles_config.get('paragraph', self.DEFAULT_STYLES_CONFIG['paragraph'])
        
        if re.match(r'h[1-6]', tag_name):
            level = int(tag_name[1])
            style_key = tag_name.upper()
            heading_style_config = self.styles_config.get(style_key, para_style_config)
            
            text_content = element.get_text(separator='').strip() 
            p = doc.add_heading(level=level) 
            
            for child_content in element.contents:
                self._add_inline_content_to_paragraph(p, child_content, heading_style_config, md_dir)
            
            # Check if paragraph is effectively empty (no text and no drawing elements like images)
            has_drawing = any(True for run in p.runs for el in run.element.xpath('.//w:drawing'))
            if not p.text.strip() and not has_drawing:
                # This heading is empty. Word usually handles empty headings fine.
                # If you want to remove it: (Caution: modifying doc.element.body can be risky)
                # try:
                #     if p._element.getparent() is not None:
                #         p._element.getparent().remove(p._element)
                # except AttributeError: # Should not happen if p is valid paragraph
                #     pass 
                pass # For now, let Word handle empty headings.

            self._apply_paragraph_format(p, heading_style_config)
            if current_list_level > 0 : 
                 p.paragraph_format.left_indent = Cm(1.27 * current_list_level) 


        elif tag_name == 'p':
            if in_list_item and not getattr(element, '_is_list_item_para_handled', False):
                p = doc.add_paragraph() # This case might need refinement based on list item logic
                element._is_list_item_para_handled = True 
            else:
                 p = doc.add_paragraph()

            self._apply_paragraph_format(p, para_style_config)
            if current_list_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * current_list_level) 

            for child_content in element.contents:
                self._add_inline_content_to_paragraph(p, child_content, para_style_config, md_dir)
            
            # Check if paragraph is effectively empty (no text and no drawing elements like images)
            has_drawing = any(True for run in p.runs for el in run.element.xpath('.//w:drawing'))
            if not p.text.strip() and not has_drawing:
                 # This paragraph is empty. Word usually handles empty paragraphs.
                 # If you want to remove it: (Caution)
                 # try:
                 #    if p._element.getparent() is not None:
                 #        p._element.getparent().remove(p._element)
                 # except AttributeError:
                 #    pass
                 pass # For now, let Word handle empty paragraphs.


        elif tag_name == 'pre':
            code_text = element.find('code').get_text() if element.find('code') else element.get_text()
            code_block_style_config = self.styles_config.get('code_block', self.DEFAULT_STYLES_CONFIG['code_block'])

            # Using a table for background color is a common workaround
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False 
            # Try to make table width responsive to nesting, assuming A4 page (approx 16cm content width)
            content_width_cm = 16.0 - (1.27 * current_list_level)
            table.columns[0].width = Cm(max(content_width_cm, 5.0)) # Minimum width 5cm

            cell = table.cell(0, 0)
            if 'background_color' in code_block_style_config:
                self._set_cell_background(cell, code_block_style_config['background_color'])
            
            # Code block paragraph within the cell
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.clear() # Clear any default run

            # Apply paragraph formatting (like line spacing) from code_block_style_config to the paragraph in cell
            self._apply_paragraph_format(p, code_block_style_config)
            # Remove first line indent for code blocks, if any was inherited
            p.paragraph_format.first_line_indent = None 

            # Add code lines with font styling
            # Split lines carefully to preserve empty lines
            lines = code_text.split('\n')
            for i, line_text in enumerate(lines):
                run = p.add_run(line_text)
                self._apply_font_style(run, code_block_style_config)
                if i < len(lines) - 1: # Add line break for all but last line
                    run.add_break()
            
            # Set space after the table itself to match paragraph spacing
            # This is done by adding an empty paragraph after the table with appropriate spacing.
            # Or, by adjusting the spacing of the paragraph containing the table (if possible, complex)
            # Easiest: Set space_after on the paragraph *inside* the cell if that's the only content.
            # The _apply_paragraph_format on 'p' already handles its internal spacing.
            # The spacing of the table element itself in the document flow is harder to control directly.
            # Add a small paragraph after the table to manage spacing if needed
            # However, the problem asks for 'space_after_pt' on the code_block style.
            # This implies the block itself should have this spacing.
            # Let's ensure the paragraph AFTER the table has the space_before from the next element,
            # and this code block's paragraph (inside cell) has its space_after.
            # This seems to be what the current _apply_paragraph_format does.


        elif tag_name in ['ul', 'ol']:
            self._process_list(doc, element, md_dir, current_list_level)

        elif tag_name == 'table':
            # This is for Markdown tables.
            # Assuming standard Markdown table structure (thead, tbody, tr, th, td)
            # Use a default table style from docx, e.g., 'TableGrid' or a custom one.
            
            html_rows = element.find_all('tr')
            if not html_rows: return

            num_cols = 0
            first_row_cells = html_rows[0].find_all(['th', 'td'])
            num_cols = len(first_row_cells)
            if num_cols == 0: return

            word_table = doc.add_table(rows=1, cols=num_cols) # Start with 1 row for header or first data row
            word_table.style = 'TableGrid' # Apply a basic grid style

            # Populate header row if <thead> or <th> exists in first row
            is_header_row = any(c.name == 'th' for c in first_row_cells)
            if is_header_row:
                header_cells_html = html_rows[0].find_all('th')
                for i, cell_html in enumerate(header_cells_html):
                    if i < num_cols:
                        cell_text = cell_html.get_text().strip()
                        # Apply basic header styling (bold, centered) to run
                        run = word_table.cell(0, i).paragraphs[0].add_run(cell_text)
                        run.bold = True
                        # TODO: Use a 'table_header' style from self.styles_config if defined
                        word_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                data_rows_html = html_rows[1:]
            else:
                # No explicit header, first row is data
                for i, cell_html in enumerate(first_row_cells):
                     if i < num_cols:
                        word_table.cell(0,i).text = cell_html.get_text().strip()
                        # TODO: Apply 'table_cell' style
                data_rows_html = html_rows[1:]

            # Populate data rows
            for html_row in data_rows_html:
                docx_row_cells = word_table.add_row().cells
                html_cells = html_row.find_all('td')
                for i, cell_html in enumerate(html_cells):
                    if i < num_cols:
                        docx_row_cells[i].text = cell_html.get_text().strip()
                        # TODO: Apply 'table_cell' style


        elif tag_name == 'hr':
            p = doc.add_paragraph()
            # Add a horizontal rule using paragraph border
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6') # Border size (e.g., 6 = 0.75pt)
            bottom_border.set(qn('w:space'), '1') # Space attribute
            bottom_border.set(qn('w:color'), 'auto') # Or a specific hex color
            pBdr.append(bottom_border)
            pPr.append(pBdr)
            # Add some spacing around HR if not handled by adjacent elements
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        elif tag_name == 'blockquote':
            # Blockquotes can contain other block elements.
            # Add a paragraph with specific styling for blockquote, then process children.
            # For Word, this might involve indenting and perhaps a left border or different font style.
            # Simplest: indent and recurse.
            # More complex: create a specific blockquote paragraph style in Word.
            # For now, just process children with increased nesting level for indentation.
            for child_element in element.children:
                if isinstance(child_element, Tag) or (isinstance(child_element, NavigableString) and child_element.strip()):
                    self._process_html_element(doc, child_element, md_dir, current_list_level + 1, in_list_item)


    def _process_list(self, doc, list_element_html, md_dir, level):
        """ Processes <ul> or <ol> list elements. `level` is the nesting depth. """
        # print(f"Processing list: {list_element_html.name}, level: {level}")
        list_items_html = list_element_html.find_all('li', recursive=False)
        
        for item_html in list_items_html:
            self._process_list_item_content(doc, item_html, md_dir, level, list_element_html.name)


    def _process_list_item_content(self, doc, item_html, md_dir, list_level, list_type_tag):
        """
        Processes the content of an <li> item.
        An <li> can contain inline text, then other block elements (like nested lists, paragraphs, code blocks).
        """
        # print(f"Processing li item, level: {list_level}")
        para_style_config = self.styles_config.get('paragraph', self.DEFAULT_STYLES_CONFIG['paragraph'])
        
        # Determine the Word list style name
        # python-docx uses 'ListBullet' or 'ListNumber' for the first level.
        # For nested lists, it automatically cycles through 'ListBullet2/ListNumber2', etc., if defined in template.
        # Or, we can manage indentation manually.
        # For simplicity and control, we'll manage indentation and use base styles.
        
        style_name = 'ListBullet' if list_type_tag == 'ul' else 'ListNumber'
        
        # The first paragraph of the list item carries the bullet/number.
        # Subsequent block elements inside <li> will be new paragraphs indented further.
        
        first_para_in_item = None
        
        # Iterate over children of <li>. The first sequence of inline content forms the main list item text.
        current_block_is_first_in_li = True

        for child_node in item_html.contents:
            if isinstance(child_node, NavigableString) and child_node.strip():
                if first_para_in_item is None:
                    first_para_in_item = doc.add_paragraph(style=style_name)
                    self._apply_paragraph_format(first_para_in_item, para_style_config) # Apply para styles
                    first_para_in_item.paragraph_format.left_indent = Cm(1.27 * list_level) # Indent based on level
                    # Negative first_line_indent for hanging bullet/number
                    first_para_in_item.paragraph_format.first_line_indent = Cm(-0.635) 

                self._add_inline_content_to_paragraph(first_para_in_item, child_node, para_style_config, md_dir)
                current_block_is_first_in_li = False

            elif isinstance(child_node, Tag):
                # Inline tags (strong, em, code, img, a)
                if child_node.name in ['strong', 'b', 'em', 'i', 'code', 'a', 'img', 'br']:
                    if first_para_in_item is None:
                        first_para_in_item = doc.add_paragraph(style=style_name)
                        self._apply_paragraph_format(first_para_in_item, para_style_config)
                        first_para_in_item.paragraph_format.left_indent = Cm(1.27 * list_level)
                        first_para_in_item.paragraph_format.first_line_indent = Cm(-0.635)
                    
                    self._add_inline_content_to_paragraph(first_para_in_item, child_node, para_style_config, md_dir)
                    current_block_is_first_in_li = False
                
                # Block-level tags (p, ul, ol, pre, table, hr, blockquote, h1-h6)
                else:
                    if first_para_in_item is None and current_block_is_first_in_li:
                        # If a block element (e.g. <p> or <pre>) is the *very first thing* in an <li>,
                        # it might need special handling or an empty leading paragraph for the bullet.
                        # For now, let's ensure there's a paragraph for the bullet if no text preceded.
                        # This typically happens if Markdown is like: `* <p>text</p>` (less common)
                        # Usually it's `* text <p>more</p>` or `* text\n  * nested`
                        # Let's assume for now that block elements start *after* some initial list item text or directly.
                        # If a block like <p> is first, _process_html_element will create its own paragraph.
                        # We need to ensure the list bullet style is applied correctly.
                        # This logic can get complex. A simpler model: list item always starts a paragraph.
                        # If the first child_node is a block, we create the list item paragraph, then process the block.
                        first_para_in_item = doc.add_paragraph(style=style_name) # Create the bullet paragraph
                        self._apply_paragraph_format(first_para_in_item, para_style_config)
                        first_para_in_item.paragraph_format.left_indent = Cm(1.27 * list_level)
                        first_para_in_item.paragraph_format.first_line_indent = Cm(-0.635)
                        # This paragraph might remain empty if the block takes all content.

                    # Process the block element. It will be indented further than the list item's bullet.
                    # The `current_list_level` for a block *inside* a list item should be `list_level + 1` for indentation purposes.
                    self._process_html_element(doc, child_node, md_dir, current_list_level=list_level + 1, in_list_item=True)
                    current_block_is_first_in_li = False # Subsequent blocks are not the first.
        
        # If li was empty or only contained whitespace, first_para_in_item might still be None.
        # Add an empty paragraph with list style to represent an empty list item.
        if first_para_in_item is None:
            first_para_in_item = doc.add_paragraph(style=style_name)
            self._apply_paragraph_format(first_para_in_item, para_style_config)
            first_para_in_item.paragraph_format.left_indent = Cm(1.27 * list_level)
            first_para_in_item.paragraph_format.first_line_indent = Cm(-0.635)



    def markdown_to_docx(self, md_file_path, docx_file_path):
        if not os.path.exists(md_file_path):
            raise FileNotFoundError(f"Markdown file '{md_file_path}' not found.")

        md_dir = os.path.dirname(os.path.abspath(md_file_path))

        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Extensions: fenced_code, tables, sane_lists are good.
        # 'nl2br' (newline to <br>) can be useful but might interact with paragraph spacing.
        # 'extra' includes many useful ones like footnotes, abbreviations, def_list, etc.
        html_content = markdown.markdown(md_content, extensions=['fenced_code', 'tables', 'sane_lists', 'extra'])
        soup = BeautifulSoup(html_content, 'html.parser')
        
        doc = Document()

        # Set default document font based on 'paragraph' style (if defined)
        # This affects the 'Normal' style, which is the base for most other styles.
        paragraph_style_config = self.styles_config.get('paragraph', self.DEFAULT_STYLES_CONFIG['paragraph'])
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = paragraph_style_config.get('font_name', 'Calibri')
        if normal_font.name:
            normal_font.element.rPr.rFonts.set(qn('w:eastAsia'), normal_font.name)
        normal_font.size = Pt(paragraph_style_config.get('font_size', 11))
        
        # Apply paragraph formatting from 'paragraph' style to 'Normal' style
        # This is less direct, usually you apply to paragraphs, not the style itself this way for all props.
        # For now, let _apply_paragraph_format handle individual paragraphs.

        # Define or update built-in styles (H1-H6, ListBullet, ListNumber)
        # This is more robust for Word interoperability.
        for i in range(1, 7):
            style_key = f'H{i}'
            if style_key in self.styles_config:
                h_style_conf = self.styles_config[style_key]
                try:
                    h_style_docx = doc.styles[f'Heading {i}']
                except KeyError: # Style might not exist if using a very blank template
                    h_style_docx = doc.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
                    h_style_docx.base_style = normal_style # Base it on Normal
                
                # Apply font
                h_font = h_style_docx.font
                h_font.name = h_style_conf.get('font_name', normal_font.name)
                if h_font.name: h_font.element.rPr.rFonts.set(qn('w:eastAsia'), h_font.name)
                h_font.size = Pt(h_style_conf.get('font_size', normal_font.size.pt))
                h_font.bold = h_style_conf.get('bold', False)
                h_font.italic = h_style_conf.get('italic', False)
                if 'color_rgb' in h_style_conf: h_font.color.rgb = RGBColor(*h_style_conf['color_rgb'])

                # Apply paragraph format to the style
                h_p_fmt = h_style_docx.paragraph_format
                if 'space_before_pt' in h_style_conf: h_p_fmt.space_before = Pt(h_style_conf['space_before_pt'])
                if 'space_after_pt' in h_style_conf: h_p_fmt.space_after = Pt(h_style_conf['space_after_pt'])
                if 'line_spacing' in h_style_conf: h_p_fmt.line_spacing = h_style_conf['line_spacing']
                # Alignment for heading styles
                alignment_map = {'LEFT': WD_ALIGN_PARAGRAPH.LEFT, 'CENTER': WD_ALIGN_PARAGRAPH.CENTER, 'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT, 'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY}
                h_p_fmt.alignment = alignment_map.get(str(h_style_conf.get('alignment', 'LEFT')).upper(), WD_ALIGN_PARAGRAPH.LEFT)


        # Process all top-level HTML elements generated by Markdown
        for element in soup.find_all(True, recursive=False): # True finds all tags
            self._process_html_element(doc, element, md_dir, current_list_level=0)
            
        doc.save(docx_file_path)
        print(f"Successfully converted '{md_file_path}' to '{docx_file_path}'")


    def load_styles(self, file_path="styles_config.json"):
        # Try to load styles, if fails or not found, self.styles_config remains default
        if os.path.exists(file_path):
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)
                
                # Update self.styles_config: start with defaults, then update with loaded
                # This ensures all necessary style keys are present.
                base_styles = {k: v.copy() for k, v in self.DEFAULT_STYLES_CONFIG.items()}
                for style_name, style_data in loaded_config.items():
                    if style_name in base_styles:
                        if isinstance(style_data, dict):
                             base_styles[style_name].update(style_data) # Update existing default style
                    else: # New style not in defaults
                        if isinstance(style_data, dict):
                            base_styles[style_name] = style_data.copy() 
                
                self.styles_config = base_styles
                self._ensure_rgb_tuples() # Convert any list RGBs to tuples
                # print(f"Styles loaded and merged from '{file_path}'.")
            except json.JSONDecodeError as e:
                print(f"Error: Could not parse styles_config.json: {e}. Using default styles.")
                self.styles_config = {k: v.copy() for k, v in self.DEFAULT_STYLES_CONFIG.items()}
            except Exception as e:
                print(f"Error loading styles: {e}. Using default styles.")
                self.styles_config = {k: v.copy() for k, v in self.DEFAULT_STYLES_CONFIG.items()}
        else:
            # print(f"Styles file '{file_path}' not found. Using default styles.")
            self.styles_config = {k: v.copy() for k, v in self.DEFAULT_STYLES_CONFIG.items()}
        self._ensure_rgb_tuples()


    def save_styles(self, file_path="styles_config.json"):
        try:
            # Prepare a serializable copy (e.g., convert RGB tuples to lists)
            serializable_config = {}
            for key, value_dict in self.styles_config.items():
                style_copy = value_dict.copy()
                if 'color_rgb' in style_copy and isinstance(style_copy['color_rgb'], tuple):
                    style_copy['color_rgb'] = list(style_copy['color_rgb'])
                serializable_config[key] = style_copy
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(serializable_config, f, ensure_ascii=False, indent=4)
            # print(f"Styles saved to '{file_path}'.")
        except Exception as e:
            print(f"Error saving styles to '{file_path}': {e}")

    def get_styles(self):
        # Return a deep copy to prevent external modification of internal state
        return {k: v.copy() for k, v in self.styles_config.items()}

    def set_styles(self, new_styles):
        # Set with a deep copy
        self.styles_config = {k: v.copy() for k, v in new_styles.items()}
        self._ensure_rgb_tuples() # Ensure RGBs are tuples after setting

    def markdown_to_html(self, md_content, md_dir=None):
        """
        Converts Markdown content to an HTML fragment string.
        Image 'src' paths are converted to absolute 'file:///' URLs if they are local.
        This HTML fragment is intended to be embedded in the QWebEngineView.
        The styles for preview are applied by WordPreviewWidget via its dynamic CSS.
        """
        html_fragment = markdown.markdown(md_content, extensions=['fenced_code', 'tables', 'sane_lists', 'extra'])
        soup = BeautifulSoup(html_fragment, 'html.parser')

        if md_dir: # md_dir is the directory of the Markdown file
            for img_tag in soup.find_all('img'):
                img_src = img_tag.get('src')
                if img_src:
                    # Skip if it's already an absolute web URL
                    if img_src.startswith(('http://', 'https://')):
                        continue
                    
                    abs_fs_path = ""
                    if img_src.startswith('file:///'):
                        # If it's already a file URL, try to normalize it
                        q_url = QUrl(img_src)
                        if q_url.isLocalFile():
                            abs_fs_path = q_url.toLocalFile()
                        else: # Malformed file URL, try to treat the rest as path
                            path_part = img_src[len('file:///'):]
                            if os.path.isabs(path_part): abs_fs_path = path_part
                            else: abs_fs_path = os.path.abspath(os.path.join(md_dir, path_part))

                    elif os.path.isabs(img_src): # Absolute local path
                        abs_fs_path = img_src
                    else: # Relative local path
                        abs_fs_path = os.path.abspath(os.path.join(md_dir, img_src))
                    
                    if abs_fs_path: # If a path was determined
                        if os.path.exists(abs_fs_path):
                            # Convert the validated absolute file system path to a file:/// URL
                            img_tag['src'] = QUrl.fromLocalFile(abs_fs_path).toString()
                        else:
                            print(f"Preview Warning: Image file not found at resolved path: {abs_fs_path} (original src: {img_src})")
                            # Optionally, could set src to a placeholder or remove it
                    # else:
                        # print(f"Preview Debug: Could not determine absolute file path for img src: {img_src}")
        
        return str(soup) # Return the modified HTML fragment